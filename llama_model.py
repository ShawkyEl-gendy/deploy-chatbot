from langchain_community.vectorstores import Chroma
import re

from utils import call_llm
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
from docx import Document as DocxDocument  # Rename to avoid conflict
import os



class LLAMA_MODEL:

    def __init__(self):
        self.conversation_history = []

    def custom_arabic_text_splitter_by_heading1(self, docx_path: str):
      doc = DocxDocument(docx_path)
      sections = []
      current_section = ""

      for para in doc.paragraphs:
          if para.style.name == 'Heading 1':
              if current_section:
                  sections.append(current_section)
                  current_section = ""
              current_section += para.text + "\n"
          else:
              current_section += para.text + "\n"

      if current_section:
          sections.append(current_section)
      return sections
    

    def process_all_docx_files(self):
        folder = os.path.dirname(os.path.abspath(__file__))
        docx_files = [f for f in os.listdir(folder) if f.lower().endswith('.docx')]

        all_sections = []
        for docx_file in docx_files:
            path = os.path.join(folder, docx_file)
            sections = self.custom_arabic_text_splitter_by_heading1(path)
            print(f"✔ {len(sections)} sections from {docx_file}")
            all_sections.extend(sections)

        return all_sections


    def get_arabic_vector_store(self):
        sections = self.process_all_docx_files()
        embeddings = HuggingFaceEmbeddings(model_name='UBC-NLP/ARBERT')
        save_to_dir = "vector_ar_DB"

        if not os.path.exists(save_to_dir):
            docs = [Document(page_content=section) for section in sections]
            docs_ids = [str(i) for i in range(len(docs))]

            vector_db = Chroma.from_documents(
                documents=docs,
                embedding=embeddings,
                persist_directory=save_to_dir,
                ids=docs_ids
            )
            print(f"✔ Vector store created with {len(docs)} documents.")
            vector_db.persist()
            print("✔ Vector store saved to disk.")

        loaded_vector_db = Chroma(
            persist_directory=save_to_dir,
            embedding_function=embeddings
        )

        return loaded_vector_db


    def prompt_template_llm(self, question, context):
        
        system_message = ".تصرف كخبير في الموضوع المطروح. أجب عن السؤال باللغة العربية فقط. إذا لم يكن لديك إجابة من المعلومات المقدمة، أخبر المستخدم بذلك و اجب من معرفتك. استخدم لغة واضحة ومبسطة. أجب بشكل مباشر ومختصر، وتجنب تضمين سؤال المستخدم في الإجابة."
        
        prompt = f"""

        السؤال: {question}
        {f"السياق: {context}" if context else ""}
        قدم إجابة واضحة ومباشرة باستخدام المعلومات المتاحة فقط. إذا كان السؤال يتطلب حسابات، قدم الحل خطوة بخطوة. إذا كانت المعطيات غير كافية، اطلب المعلومات الإضافية بوضوح من المستخدم. استخدم اللغة العربية الفصحى والمبسطة دون إدخال حروف أجنبية.
        """

        prompt_template = ChatPromptTemplate.from_messages(
            [
                ("system", system_message),
                MessagesPlaceholder(variable_name="chat_history"),
                ("human", "{input}")
            ]
        )

        final_prompt = prompt_template.format(input=prompt, chat_history=self.conversation_history)
        response = call_llm(final_prompt)

        self.conversation_history.append(HumanMessage(content=question))
        self.conversation_history.append(AIMessage(content=response))

        return response

    def get_retriever(self):
        return self.get_arabic_vector_store().as_retriever(search_kwargs={"k": 3})

    def rag_chain(self, question):
        retrieved_docs = self.get_retriever().invoke(question)
        formatted_context = " ".join(doc.page_content.replace('\n', ' ') for doc in retrieved_docs)
        return self.prompt_template_llm(question, formatted_context)


    def get_important_facts(self, question):
        response = self.rag_chain(question)
        if not response.strip() or response =="":
            return ""
        return response

    def reset_conversation(self):
        self.conversation_history = []
